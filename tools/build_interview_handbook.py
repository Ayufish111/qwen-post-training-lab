from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PROJECT_DIR = Path(__file__).resolve().parents[1]
OUTPUT_PATH = PROJECT_DIR / "Qwen3后训练项目_面试深挖与代码阅读手册.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(36, 44, 53)
MUTED = RGBColor(95, 105, 118)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
PALE_YELLOW = "FFF7E0"
WHITE = RGBColor(255, 255, 255)


def set_run_font(run, name="Microsoft YaHei", size=10.5, bold=None, color=INK):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    if bold is not None:
        run.bold = bold


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_widths(table, widths_inches):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        for index, width in enumerate(widths_inches):
            row.cells[index].width = Inches(width)
            set_cell_margins(row.cells[index])
            row.cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_run_font(run, size=8.5, color=MUTED)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr, fld_char2])
    run2 = paragraph.add_run(" 页")
    set_run_font(run2, size=8.5, color=MUTED)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.78)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.82)
    section.right_margin = Inches(0.82)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.22

    heading_specs = {
        "Heading 1": (16, BLUE, 16, 8),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (11.5, DARK_BLUE, 8, 4),
    }
    for style_name, (size, color, before, after) in heading_specs.items():
        style = styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    code_style = styles.add_style("CodeBlock", WD_STYLE_TYPE.PARAGRAPH)
    code_style.font.name = "Consolas"
    code_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    code_style.font.size = Pt(7.8)
    code_style.font.color.rgb = RGBColor(35, 45, 55)
    code_style.paragraph_format.left_indent = Inches(0.16)
    code_style.paragraph_format.right_indent = Inches(0.08)
    code_style.paragraph_format.space_before = Pt(2)
    code_style.paragraph_format.space_after = Pt(5)
    code_style.paragraph_format.line_spacing = 1.0

    qa_style = styles.add_style("QAAnswer", WD_STYLE_TYPE.PARAGRAPH)
    qa_style.font.name = "Microsoft YaHei"
    qa_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    qa_style.font.size = Pt(10)
    qa_style.paragraph_format.left_indent = Inches(0.18)
    qa_style.paragraph_format.space_after = Pt(5)
    qa_style.paragraph_format.line_spacing = 1.18

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run_font(header.add_run("Qwen3 中文后训练 | 面试代码手册"), size=8.5, color=MUTED)
    add_page_number(section.footer.paragraphs[0])


def add_title_page(doc):
    for _ in range(5):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    set_run_font(p.add_run("Qwen3 中文后训练项目"), size=26, bold=True, color=DARK_BLUE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(22)
    set_run_font(p.add_run("面试深挖与代码阅读手册"), size=20, bold=True, color=BLUE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    set_run_font(p.add_run("数据治理 · QLoRA SFT · 消融实验 · Multi-IF 评测"), size=11.5, color=MUTED)
    doc.add_paragraph()
    callout = doc.add_table(rows=1, cols=1)
    set_table_widths(callout, [6.5])
    shade_cell(callout.cell(0, 0), LIGHT_BLUE)
    p = callout.cell(0, 0).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    set_run_font(
        p.add_run("用途：从项目总述进入源码，准备 30 分钟以上的算法实习项目深挖"),
        size=10.5,
        bold=True,
        color=DARK_BLUE,
    )
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run("基于仓库当前代码生成 · 2026-08-03"), size=9, color=MUTED)
    doc.add_page_break()


def add_contents(doc):
    doc.add_heading("目录", level=1)
    entries = [
        ("第1-2章", "使用说明、项目地图、冻结实验矩阵与当前证据"),
        ("第3章A", "configs/project.yaml：统一配置"),
        ("第3章B", "build_sft_datasets.py：数据构建与清洗"),
        ("第3章C-E", "Schema、自动验证与人工审核"),
        ("第3章F", "preprocess_sft.py：模板、token与labels"),
        ("第3章G", "train_sft.py：QLoRA训练"),
        ("第3章H-J", "评测集冻结、推理和Multi-IF正式评测"),
        ("第4章", "30分钟以上深挖题库：标准回答与继续追问"),
        ("第5章", "真实面试节奏、三分钟开场稿与GitHub检查"),
        ("附录", "文件职责速查与最终复习标准"),
    ]
    for label, detail in entries:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.2)
        p.paragraph_format.space_after = Pt(5)
        set_run_font(p.add_run(f"{label}  "), size=10.5, bold=True, color=DARK_BLUE)
        set_run_font(p.add_run(detail), size=10.2)
    add_note(doc, "导航", "所有章节均使用Word标题样式，打开“导航窗格”即可按标题跳转。", LIGHT_BLUE)
    doc.add_page_break()


def add_note(doc, label, text, fill=PALE_YELLOW):
    table = doc.add_table(rows=1, cols=1)
    set_table_widths(table, [6.5])
    shade_cell(table.cell(0, 0), fill)
    p = table.cell(0, 0).paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    set_run_font(p.add_run(f"{label}："), bold=True, color=DARK_BLUE)
    set_run_font(p.add_run(text), size=10)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.42)
        p.paragraph_format.first_line_indent = Inches(-0.2)
        p.paragraph_format.space_after = Pt(3)
        set_run_font(p.add_run(item), size=10.2)


def add_code(doc, path, start, end):
    lines = (PROJECT_DIR / path).read_text(encoding="utf-8").splitlines()
    start = max(1, start)
    end = min(len(lines), end)
    text = "\n".join(f"{number:>4} | {lines[number - 1]}" for number in range(start, end + 1))
    p = doc.add_paragraph(style="CodeBlock")
    p.paragraph_format.keep_together = False
    set_run_font(p.add_run(text), name="Consolas", size=7.8, color=RGBColor(35, 45, 55))
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F5F7F9")
    p_pr.append(shd)


def add_explanation(doc, title, text, why=None, interview=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    set_run_font(p.add_run(title), size=10.5, bold=True, color=DARK_BLUE)
    set_run_font(p.add_run(" " + text), size=10)
    if why:
        p = doc.add_paragraph(style="QAAnswer")
        set_run_font(p.add_run("为什么这样做："), size=9.8, bold=True, color=BLUE)
        set_run_font(p.add_run(why), size=9.8)
    if interview:
        p = doc.add_paragraph(style="QAAnswer")
        set_run_font(p.add_run("面试抓手："), size=9.8, bold=True, color=RGBColor(122, 90, 0))
        set_run_font(p.add_run(interview), size=9.8)


def add_file_chapter(doc, title, path, purpose, blocks):
    doc.add_page_break()
    doc.add_heading(title, level=1)
    p = doc.add_paragraph()
    set_run_font(p.add_run("文件："), bold=True, color=DARK_BLUE)
    set_run_font(p.add_run(path), name="Consolas", size=9.5)
    p = doc.add_paragraph()
    set_run_font(p.add_run("职责："), bold=True, color=DARK_BLUE)
    set_run_font(p.add_run(purpose), size=10.2)
    add_note(
        doc,
        "阅读方法",
        "以下按最小逻辑块解释。连续括号、纯格式换行和重复字段并入相邻语句，但所有会改变数据、训练或评测行为的代码均覆盖。",
        LIGHT_BLUE,
    )
    for heading, start, end, explanation, why, interview in blocks:
        doc.add_heading(f"{heading}（L{start}-L{end}）", level=2)
        add_code(doc, path, start, end)
        add_explanation(doc, "逐行读法：", explanation, why, interview)


def add_overview(doc):
    doc.add_heading("使用说明与项目地图", level=1)
    add_note(
        doc,
        "先后顺序",
        "先读第1-2章建立主线，再按第3章的文件顺序看代码；最后用第4章问题库做闭卷复述。不要一开始死记 TrainingArguments。",
    )
    doc.add_heading("1. 一句话项目定义", level=2)
    doc.add_paragraph(
        "基于 Qwen3-4B-Base 构建可审计的中文后训练流程：从 Infinity-Instruct 7M 中形成冻结数据，使用 QLoRA 完成 SFT，通过控制变量实验分别研究 LoRA 挂载范围、数据质量和数据规模，并用 Multi-IF 中文多轮基准验证指令遵循能力。"
    )
    doc.add_heading("2. 数据与模型的完整流向", level=2)
    p = doc.add_paragraph(style="CodeBlock")
    flow = (
        "Infinity-Instruct 7M\n"
        "  -> 技术过滤/格式转换/去重/风险评分\n"
        "  -> full_clean_10000 + ablation_clean_2000 + ablation_raw_2000\n"
        "  -> 95/5 split + Qwen chat template + labels mask + 长度门禁\n"
        "  -> QLoRA: S1 attention / S2 all-linear / S3 raw / S4 full\n"
        "  -> Multi-IF Chinese 454 x 最多3轮\n"
        "  -> 消融表、错误分析、最终结论"
    )
    set_run_font(p.add_run(flow), name="Consolas", size=9.5)

    doc.add_heading("3. 冻结实验矩阵", level=2)
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    set_table_widths(table, [0.65, 1.55, 1.55, 1.15, 1.6])
    headers = ["ID", "数据", "LoRA target", "训练预算", "研究问题"]
    for i, text in enumerate(headers):
        shade_cell(table.rows[0].cells[i], LIGHT_BLUE)
        p = table.rows[0].cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run_font(p.add_run(text), size=9, bold=True, color=DARK_BLUE)
    rows = [
        ("B0", "无", "无", "不训练", "Base基线"),
        ("S1", "clean 2k", "q/k/v/o", "200 steps", "注意力方案"),
        ("S2", "同一clean 2k", "all-linear", "200 steps", "挂载范围消融"),
        ("S3", "raw 2k", "S1/S2胜者", "200 steps", "数据质量消融"),
        ("S4", "full clean 10k", "S1/S2胜者", "2 epochs", "最终SFT/规模"),
    ]
    for row_data in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row_data):
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i in {0, 3} else WD_ALIGN_PARAGRAPH.LEFT
            set_run_font(p.add_run(value), size=8.8)

    doc.add_heading("4. 当前已经取得的证据", level=2)
    add_bullets(
        doc,
        [
            "冻结数据：full clean 10,000；clean/raw 消融各2,000；clean是full子集，raw与两者不重叠。",
            "人工分层审核：100条中99条pass、1条major；这只是审计，不代表逐条事实核验。",
            "token化后：clean/raw训练各1,815、验证各93；full训练9,249、验证485；配对超长移除92组。",
            "B0 Multi-IF official average：turn1 40.45%，turn2 20.46%，turn3 16.07%。",
            "S1：train loss 1.3296、eval loss 1.4947、约19.8分钟；S2：1.3057、1.4935、约24.1分钟。",
            "S1/S2正式Multi-IF仍以最终summary为准，不能凭接近的eval loss选胜者。",
        ],
    )


def file_guides():
    return [
        (
            "第3章A：统一配置",
            "configs/project.yaml",
            "集中冻结模型、数据、QLoRA、训练和评测参数，避免脚本之间悄悄漂移。",
            [
                ("项目与模型", 1, 12, "第1-2行给项目命名并固定随机种子。model段指定Base模型、revision、是否信任远程代码和最大序列长度。revision必须记录，否则上游模型更新会破坏复现。", "所有训练与评测都从同一配置读取模型身份。", "追问：为什么Base而不是Instruct？答：需要观察自己的SFT带来的增量，Instruct已有未知后训练。"),
                ("数据冻结参数", 14, 19, "记录上游数据源、子集、full/ablation目标规模、划分比例与数据门禁状态。approved是训练前的人工闸门，不是模型自动判断。", "把数据治理状态和训练解耦，防止未审核数据误启动昂贵训练。", "追问：approved能证明每条正确吗？不能，只代表预定义门禁通过。"),
                ("QLoRA参数", 21, 27, "开启4-bit NF4、double quant；LoRA rank为16、alpha为32、dropout为0.05。alpha/r=2是LoRA更新的缩放。", "在24GB显卡上降低基座权重和优化器显存，同时保留可训练低秩适配器。", "追问：LoRA矩阵是否也是4-bit？通常不是，LoRA参数以bf16/fp16训练。"),
                ("公共训练参数", 29, 38, "batch=2，梯度累积8次形成有效batch 16；学习率2e-4；linear scheduler；warmup 5%；logging每10步。", "将所有实验公共超参数冻结，控制变量只允许来自实验矩阵。", "追问：梯度累积与真实batch完全等价吗？在无batch相关随机差异时梯度近似等价，但吞吐、数值舍入和调度粒度不同。"),
                ("S1-S4实验定义", 40, 62, "S1/S2固定数据、只改target；S3/S4的target为空，要求命令行传入S1/S2胜者。max_steps=200会覆盖epoch；S4用-1表示按2个epoch完整训练。", "用配置结构强制实验协议，而不是靠人记住。", "追问：S4为什么不从S1继续？因为那会同时改变初始化和训练步数，无法归因于数据规模。"),
                ("评测冻结参数", 64, 69, "指定B0、Multi-IF中文、pilot=100、最大生成512、关闭采样。", "所有模型用完全相同的解码规则，否则指标不可比较。", "追问：关闭采样的代价？降低输出多样性，但显著提升评测复现性。"),
            ],
        ),
        (
            "第3章B：数据构建与清洗",
            "src/build_sft_datasets.py",
            "从上游行转换、流式抽样、硬错误过滤、软风险评分，再构造full/clean/raw三份冻结数据。",
            [
                ("导入、路径与规则常量", 1, 44, "导入哈希、JSON、随机数、正则、统计和分层容器；固定seed、数据规模、缓存路径。TASK_RULES用能力标签和问题关键词映射任务桶；placeholder正则和模型自称词用于可观察风险检查。", "把规则写成显式常量，便于审计和复跑。", "追问：关键词分类会误判吗？会，所以它只用于粗粒度分层，不作为语义标签真值。"),
                ("读写、标准化与稳定哈希", 46, 69, "read/write负责JSONL；normalized做NFKC、小写和仅保留字母数字；question_key据此去重；stable_number用BLAKE2b生成稳定排序值。", "NFKC统一全角半角等兼容字符；稳定哈希避免Python内置hash跨进程随机化。", "追问：这能抓住语义重复吗？不能，只是规范化后的精确重复。"),
                ("上游行转换", 71, 108, "依次验证中文、恰好两轮、human/gpt角色、字符串非空、reward数值；随后映射到messages和metadata。任何一项失败返回None。", "先做技术合法性过滤，避免后续代码处理异构或损坏schema。", "追问：为什么只留一问一答？当前SFT目标是单轮监督，保持监督结构简单；代价是无法直接训练多轮保持能力。"),
                ("流式蓄水池抽样", 110, 140, "ModelScope以streaming读取7M。前30k合格样本直接入池；第n个合格样本以30k/n概率替换池中元素。这样不把全部数据放进内存，每个合格样本入池概率相同。", "在大规模流式数据上获得近似均匀样本，内存复杂度O(30k)。", "追问：为什么不能只取前30k？上游可能按来源排序，前缀会产生严重分布偏差。"),
                ("缓存优先级", 142, 150, "优先使用technical_candidates缓存；若不存在但冻结full/raw已在本地，就从12k实验宇宙确定性重建；都不存在才重新扫描7M。", "避免每次实验重新下载并扫描数百万行。", "追问：这意味着什么复现边界？从12k重建只能复现当前冻结实验，不能重新证明最初30k候选如何产生。"),
                ("任务桶与长度桶", 152, 169, "task_bucket拼接能力标签和问题文本，命中规则后归类，否则general；length_bucket按问答字符总长分成<=256/512/1024/2048和更长。", "为clean/raw匹配可观察的任务结构和长度成本。", "追问：为什么不是token长度？构建阶段不加载模型tokenizer以降低成本，最终仍在预处理阶段做token门禁。"),
                ("硬错误、软风险和质量分", 171, 230, "先检查messages/metadata结构，再检查空文本、问答相同、乱码和占位符。reward低、答案过短、样本过长、自称模型、复制问题只打flag。分数=50+裁剪reward×2-风险惩罚。", "客观结构错误可以硬删；事实/专业正确性不应由简单规则武断删除。", "追问：惩罚权重有理论最优吗？没有，是透明启发式；论文级研究需敏感性分析或学习式评分。"),
                ("去重与元数据标注", 232, 257, "逐行调用inspect_row；硬错误计数后删除；question_key进入seen实现首次保留；随后把score、flags、task/length bucket写回metadata。", "训练文本与审计信息分离，既保持messages标准，又保留可追溯证据。", "追问：为什么按问题去重而不是问答对？相同问题的多个答案会导致监督冲突；但也会丢失可能的多样高质量答案。"),
                ("精确分层配额", 259, 280, "先按源分布计算每层理想数量，再在minimum和capacity约束内取整；循环补足或削减直到目标总数严格等于2k/10k。", "简单按比例取整会出现总数不等于目标或小桶完全消失。", "追问：minimum=1的意义？保证每个有效分层至少有代表；代价是极小桶可能被相对过采样。"),
                ("构造full/clean/raw", 282, 327, "按(task,length)分组并按质量分降序、稳定哈希打破平局。每层顶部进入full，顶部对应数量进入clean，尾部进入raw；最后断言规模严格为10k/2k/2k。", "clean/raw在任务和长度层面匹配，但质量分明显分离。", "追问：这是严格配对吗？不是同题配对，只是分层计数匹配，隐藏难度仍可能不同。"),
                ("审计报告与主入口", 329, 469, "计算分布、均值中位数、SHA256并写报告；main执行加载、标注、选择、集合关系断言、写出文件。断言失败立即停止，避免静默产生错误实验集。", "把数据结果变成可复核工件，而不是只在控制台打印一次。", "追问：SHA256有什么用？锁定具体文件字节，确认云端和本地使用的是同一版本。"),
            ],
        ),
        (
            "第3章C：数据Schema",
            "src/data_schema.py",
            "为SFT和未来DPO记录提供可复用的结构验证。",
            [
                ("角色集合与消息验证", 1, 28, "VALID_ROLES限制允许角色。validate_messages拒绝字符串伪装成序列，逐项检查消息对象、role以及content/tool_calls，错误全部收集而非遇到第一项就退出。", "一次返回完整错误列表，便于批量数据审计。", "追问：为什么允许tool？为后续Agent/DPO格式预留，但当前SFT数据仍只用user/assistant。"),
                ("SFT记录验证", 30, 39, "验证record是Mapping、messages合法，并强制最后一条是assistant，因为监督目标必须存在。", "避免最后一条为user时产生没有目标答案的SFT样本。", "追问：system消息是否允许？允许，但当前清洗器生成的固定两轮数据没有system。"),
                ("DPO记录验证", 41, 48, "要求prompt/chosen/rejected三个字段都存在且各自是合法消息列表。", "DPO需要同一prompt下的偏好答案对。", "追问：这里只有schema，没有验证chosen一定优于rejected；偏好质量仍需单独门禁。"),
            ],
        ),
        (
            "第3章D：冻结数据验证",
            "scripts/validate_dataset.py",
            "在训练前后分别检查原始JSONL和tokenized cache，形成自动质量门禁。",
            [
                ("文件与标准化定义", 1, 43, "固定三份数据的路径和期望行数；question_key使用与清洗一致的NFKC逻辑；stratum返回任务桶和长度桶。", "验证器必须独立重算关键属性，不能只相信构建脚本输出。", "追问：若清洗器和验证器共享同一个bug怎么办？独立实现能降低但不能消除共同假设风险。"),
                ("逐文件schema与重复检查", 45, 99, "读取每份JSONL，逐条运行validate_sft_record，检查必要metadata字段、行数和重复数，并打印score/reward统计。", "同时验证格式正确性与实验规模，防止缺行仍继续训练。", "追问：为什么只展示前三个错误？防止日志爆炸，同时最终保留错误总数。"),
                ("集合关系与分层匹配", 101, 130, "计算full/clean/raw问题集合，验证clean属于full、raw不与full/clean重叠、clean/raw层计数完全一致。", "这些是消融实验成立的前提，而不是普通数据质量检查。", "追问：为什么raw不能与full重叠？否则S4与S3的数据边界不清晰，结果解释更复杂。"),
                ("tokenized门禁", 133, 177, "若cache存在则检查六个split、三张量等长、长度<=1024、至少一个有效label、有效label等于对应input_id，并检查clean/raw划分大小一致。", "直接验证Trainer真正吃到的数据，而不只验证源JSON。", "追问：还缺什么？没有在这里逐样本验证prompt位置全是-100，可进一步增强。"),
                ("失败即退出", 180, 184, "聚合raw和tokenized失败；存在任意失败就打印并以状态1退出，否则输出VALIDATION PASSED。", "便于CI或训练脚本把它作为真正门禁。", "追问：为什么不是warning？核心实验不变量一旦破坏就不能继续。"),
            ],
        ),
        (
            "第3章E：人工分层审核",
            "scripts/audit_sft.py",
            "对full clean执行可恢复、可复现的100条人工审核，并把结果写入审计报告。",
            [
                ("路径与稳定样本键", 1, 29, "固定seed、样本数和进度/报告路径。sample_key以问题SHA256识别样本。", "即使JSONL行顺序变化，同一个问题仍有稳定身份。", "追问：只哈希问题有什么风险？相同问题不同答案会共享键；但上游已按问题去重。"),
                ("分层抽样", 31, 49, "按task_bucket分组，每桶先稳定抽一条，再从剩余数据seed shuffle补齐100，最后再次固定洗牌。", "避免小任务桶在纯随机抽样中完全缺席。", "追问：这是不是比例分层？不是，先保证覆盖再随机填充，整体比例近似但不严格。"),
                ("进度保存与报告", 52, 94, "每次判断写入隐藏进度文件；完成后统计pass/minor/major和任务覆盖，将非pass样本及原因插入报告标记区。", "长时间人工审核可中断恢复，报告更新不会重复追加。", "追问：为什么报告只列问题前120字符？控制可读性，原始数据仍保留完整文本。"),
                ("交互审核主循环", 97, 157, "跳过已完成样本，显示任务、质量分、风险、问题和回答；1/2/3映射pass/minor/major；非pass要求理由；q保存退出。", "人工判断用于发现规则抓不到的语义错误。", "追问：审核者偏差如何控制？当前仅单人审核，论文级工作需双人标注和一致性指标。"),
            ],
        ),
        (
            "第3章F：切分、模板与Labels",
            "src/preprocess_sft.py",
            "把三份冻结文本转成Trainer直接消费的六个tokenized split。",
            [
                ("配置、tokenizer与读取", 1, 39, "从统一YAML取得模型、revision、max length和seed，确定cache目录并加载Qwen tokenizer。read_jsonl只负责读取冻结文件。", "预处理必须使用与训练相同的tokenizer和revision。", "追问：ModelScope下载的数据为什么仍能用HF tokenizer接口？ModelScope是数据/模型分发来源，模型目录和Transformers接口是两个概念。"),
                ("full与prompt两次模板化", 41, 61, "full包含user+assistant且不补生成前缀；prompt使用messages[:-1]并补assistant开始标记。随后验证prompt_ids确实是full的前缀。", "精确定位监督答案起点，避免用字符串长度猜token边界。", "追问：add_generation_prompt会生成回答吗？不会，只补模板中的assistant起始token。"),
                ("构造labels", 63, 70, "input_ids取完整对话；prompt长度范围全部设-100，后续答案token直接复制input_ids，包括im_end。", "问题提供条件但不贡献loss；模型学习答案和停止位置。", "追问：为什么不能把im_end也mask？模型会更难学会在合适位置停止。"),
                ("确定性95/5切分", 73, 82, "生成索引、seed shuffle、round 5%且至少1条验证；Dataset.map返回新Dataset，并remove原始列。", "用索引使clean/raw可以共享同一划分位置。", "追问：datasets.map是否原地修改？不是，必须接收返回值。"),
                ("配对长度过滤", 85, 102, "clean/raw先分别tokenize；只有同一索引两边都<=1024才保留，然后把共享有效索引应用到train和validation。", "避免长度过滤后两组样本量或分层位置不同。", "追问：为什么不truncation=True？可能截断答案并产生错误监督，过滤更可解释。"),
                ("full过滤与保存", 104, 129, "full独立切分和tokenize，再过滤超长；最终组装六个split、save_to_disk并打印规模。", "磁盘cache避免每次训练重复模板化和tokenization。", "追问：save_to_disk保存什么？Arrow数据和DatasetDict元信息，不是JSONL。"),
            ],
        ),
        (
            "第3章G：QLoRA训练",
            "src/train_sft.py",
            "读取冻结实验定义，加载4-bit基座、挂LoRA、训练、验证并保存adapter。",
            [
                ("导入、配置与target策略", 1, 43, "导入PEFT、Transformers和checkpoint工具；从YAML取得模型、cache和输出路径；attention明确列q/k/v/o，all-linear由PEFT识别全部线性层。", "把可变的target策略集中定义，防止不同实验脚本复制漂移。", "追问：all-linear会不会包括lm_head？PEFT的all-linear通常排除输出头；最终应以print_trainable_parameters和模块清单验证。"),
                ("命令行和实验门禁", 45, 99, "要求experiment；S3/S4必须显式传胜者；固定策略实验禁止传入冲突值；检测已有checkpoint，防止误覆盖，只有--resume才能继续。", "把实验协议编码成错误检查。", "追问：为什么不能自动选eval loss最低者？选择依据是外部指令评测，不是训练目标本身。"),
                ("数据审批与加载", 101, 132, "approved不通过立即终止；加载tokenizer、设置右padding和pad token；加载cache并检查所需split存在。", "训练前验证输入契约，尽早失败。", "追问：把eos当pad是否会混淆？attention_mask屏蔽pad；labels的pad再由collator置-100。"),
                ("GPU、4-bit与dtype", 134, 157, "要求CUDA；支持时用bf16，否则fp16；BitsAndBytes配置NF4、double quant和计算dtype；基座以4-bit加载到GPU0。", "4-bit降低冻结基座权重显存，bf16提高计算稳定性。", "追问：NF4是什么？针对近似正态分布权重设计的非均匀4-bit码本。"),
                ("k-bit准备与LoRA", 159, 173, "关闭use_cache以兼容checkpointing；prepare_model_for_kbit_training处理冻结、输入梯度等；LoraConfig定义因果LM、target、r/alpha/dropout和无bias；get_peft_model注入适配器。", "只训练低秩增量，原始W保持冻结。", "追问：公式W'=W+(alpha/r)BA；rank越高容量和参数越大。"),
                ("评测/保存调度", 175, 191, "pilot每50步eval/save；S4按epoch；注释说明causal LM直接使用内置eval_loss，不需要分类compute_metrics。", "pilot需要较密观察，full按epoch降低验证开销。", "追问：为什么eval loss不是最终指标？它衡量teacher-forcing token预测，不直接等于自由生成的约束遵循。"),
                ("TrainingArguments", 193, 218, "设置输出、步数/epoch、batch、累积、学习率、调度、warmup、8-bit优化器、bf16/fp16、checkpointing、日志、最佳模型和随机种子。max_steps>0覆盖epoch。", "冻结公共参数确保消融只改变目标变量。", "追问：paged AdamW 8-bit节省优化器状态显存；gradient checkpointing用额外前向重算换激活显存。"),
                ("动态padding与Trainer", 220, 237, "DataCollatorForSeq2Seq按batch动态padding，pad到8倍数，labels补-100；Trainer接收模型、参数、训练/验证集、collator和tokenizer。", "避免全数据padding到1024造成算力浪费，并利用Tensor Core友好形状。", "追问：为什么用Seq2Seq collator训练causal LM？它能正确分别padding input和labels；核心在提供的数据契约。"),
                ("日志、训练、评估与保存", 238, 271, "打印实验身份、数据量、有效batch、训练预算和可训练参数；trainer.train支持checkpoint恢复；分别保存train/eval指标、state、最终adapter和tokenizer。", "adapter是增量权重，推理时仍需基座模型。", "追问：final_adapter和checkpoint区别？checkpoint还含优化器/调度/RNG用于续训，final_adapter用于部署推理。"),
            ],
        ),
        (
            "第3章H：冻结Multi-IF中文集",
            "scripts/prepare_multi_if.py",
            "下载官方Multi-IF、筛出中文、检查精确训练重叠并写成冻结CSV。",
            [
                ("规范化和读取提示", 1, 33, "定义官方数据名与路径；normalized复用NFKC逻辑；prompt_content兼容字符串JSON或对象；空轮次返回None。", "统一训练/评测文本比较方式。", "追问：CSV字段为什么含JSON字符串？官方数据每个turn的message和约束kwargs是结构化字段。"),
                ("训练问题集合", 35, 42, "从full clean和raw读取所有训练问题的规范化key。clean已包含于full，无需重复读取。", "用于检测评测提示与训练提示的精确污染。", "追问：这里只能发现精确重叠，不能发现改写后的污染。"),
                ("中文筛选和确定顺序", 45, 57, "下载全集、统计语言、只保留Chinese；按key的SHA256排序，冻结稳定的pilot前缀。", "--limit 100每次会选同一批题，不依赖上游行顺序。", "追问：为什么不是random.seed shuffle？哈希排序更容易跨环境稳定。"),
                ("污染检查和冻结", 59, 81, "逐行逐turn检查与训练问题精确重叠；发现即报错；随后按官方字段写UTF-8 BOM CSV、计算SHA256并打印。", "固定评测数据和文件哈希，保证B0/S1-S4比较一致。", "追问：UTF-8-sig的作用？兼容Windows/Excel读取中文CSV。"),
            ],
        ),
        (
            "第3章I：通用推理与冒烟评测",
            "src/evaluate_instruction.py",
            "提供4-bit基座/adapter加载、Qwen生成函数和20条本地规则冒烟测试。",
            [
                ("路径、读取和文件哈希", 1, 39, "加载配置，定义默认冒烟数据/输出目录；read_jsonl读取样本；sha256_file分块计算，避免一次加载大文件。", "记录每次评测的输入身份。", "追问：冒烟集为什么不能作为正式结论？样本少且是本地自定义规则，覆盖和独立性不足。"),
                ("本地规则检查器", 41, 73, "支持exact、包含/排除、行数、编号列表、项目符号、JSON、字符范围和禁止正则；未知类型直接报错。", "快速发现模板、加载或明显指令遵循回归。", "追问：规则检查的优势是确定、免费；缺点是不能判断开放回答语义质量。"),
                ("4-bit加载Base和adapter", 76, 115, "要求CUDA，重置峰值显存；选择bf16/fp16；构造4-bit配置；加载tokenizer和Base；若传adapter则PeftModel挂载；eval模式并同步计时。", "B0和SFT模型共享同一加载函数，减少评测链路差异。", "追问：adapter不是完整模型，PeftModel将低秩增量应用到Base对应模块。"),
                ("生成函数", 118, 153, "apply_chat_template只生成文本模板；tokenizer转tensor；输入移动到嵌入层设备；inference_mode生成；切掉输入前缀，只解码新token，并记录token数和时间。", "避免把prompt也当成回答评分。", "追问：为什么add_special_tokens=False？chat template已经插入Qwen控制token，再添加可能重复。"),
                ("冒烟主循环", 155, 219, "解析参数、明确警告非正式benchmark；加载样本和模型；逐题生成、执行全部checks，所有规则通过才算题目pass；边生成边flush写JSONL。", "中途崩溃时尽量保留已生成证据。", "追问：这里没有resume，所以正式长评测使用evaluate_multi_if的resume。"),
                ("汇总可复现信息", 221, 285, "计算通过率、分类通过率、运行耗时、tokens/s和峰值显存；记录模型、revision、量化、dtype、版本、GPU、数据哈希并写summary。", "结果不仅有分数，还有复现实验环境。", "追问：还应记录什么？adapter哈希、git commit和依赖锁文件。"),
            ],
        ),
        (
            "第3章J：Multi-IF正式评测",
            "src/evaluate_multi_if.py",
            "调用官方约束类，对中文多轮生成计算strict/loose、prompt/instruction四类指标。",
            [
                ("结构字段解析", 1, 49, "导入通用生成函数；parse_json兼容字符串和已解析对象；parse_message处理空轮；parse_constraints读取每轮instruction id和kwargs，并验证数量一致。", "官方CSV包含嵌套JSON，必须恢复结构后实例化规则。", "追问：为什么不能用正则直接解析CSV字段？嵌套JSON转义复杂，必须用结构化解析。"),
                ("加载官方规则与版本", 51, 71, "检查third_party/Multi-IF/ifeval.py存在，将目录放入sys.path后动态import；尝试读取官方规则仓库git revision。", "使用官方判分逻辑，而不是自行近似实现。", "追问：当前revision为null是复现缺口，应保留git目录或手工记录commit。"),
                ("strict与loose判定", 74, 96, "loose生成去首行、去末行、去两端和去星号等候选；每个instruction class按kwargs构建约束并调用check_following；任一候选通过即记True。", "区分核心约束失败与轻微格式包裹。", "追问：loose不是语义评分，只是官方允许的格式容错。"),
                ("四类指标", 98, 149, "逐turn累计prompt数、instruction数和四类正确数。prompt要求该题所有约束都通过；instruction逐约束计分；official overall是四项算术平均。", "prompt指标更严格，能揭示多约束组合失败；instruction显示局部能力。", "追问：为什么prompt通常更低？约束越多，全通过概率会乘法下降。"),
                ("参数、数据和模型加载", 151, 189, "解析experiment/model/adapter/data/repo/output/limit/max tokens/resume；先加载官方规则和CSV，再调用共享4-bit加载函数。", "保证B0与adapter只在adapter参数上不同。", "追问：limit用于pilot，正式报告必须为454且limit=null。"),
                ("断点续跑", 190, 207, "--resume时读取现有JSONL，忽略空行和末尾不完整JSON，收集completed_ids；以追加模式打开文件。", "长评测断线后不重复花费已经完成的生成。", "追问：为什么按row id而不是行号？数据顺序变化时id仍稳定。"),
                ("三轮对话生成与评分", 208, 255, "逐row跳过已完成id；每轮把user消息追加到messages，生成response，按该轮约束算strict/loose，再把assistant回答放回上下文。每完成一题立即flush。", "真实模拟多轮：上一轮回答会影响后续轮次，错误能够传播。", "追问：这也是turn2/3下降的原因之一，同时后续约束数量更多。"),
                ("summary与输出", 257, 290, "记录官方数据/规则、实验、模型、adapter、行数、哈希、解码、加载时间、峰值显存和各轮指标，写answers JSONL与summary JSON。", "原始回答支持错误分析，summary支持表格比较。", "追问：只保留summary为什么不够？无法检查奖励漏洞、截断、语言异常和规则误判。"),
            ],
        ),
    ]


QUESTIONS = [
    ("项目定位", "为什么使用Qwen3-4B-Base而不是Instruct？", "Base没有叠加未知指令微调，我能把B0到SFT的变化归因于自己的数据和训练。代价是Base初始对话能力更弱，chat template兼容性也必须验证。", "如果换Instruct，需要把结论改为领域继续微调，而不是从Base构建指令能力。"),
    ("项目定位", "你的项目究竟是中文领域微调还是通用指令微调？", "数据覆盖代码、知识、数学、写作等多个任务，因此不是单一垂直领域；更准确的目标是中文通用指令遵循后训练。Multi-IF主要验证约束遵循，不代表全部领域能力。", "不要把“中文”误说成一个业务领域。"),
    ("数据", "reward为什么出现负数？", "reward是上游数据的质量/偏好信号，零点和尺度由上游打分机制决定，负数表示相对较差而不是数学上的非法数据。因为它不是我的事实裁判，所以只作为软信号进入质量分。", "若面试官追问来源，要承认未重建上游reward模型，不能杜撰。"),
    ("数据", "为什么不直接reward>0？", "会把上游打分偏差当真值，并可能删除有用的困难任务。我的做法是裁剪reward影响，再结合可观察风险排序，同时用人工抽样估计残余错误。", "硬删只针对能被确定观察到的结构故障。"),
    ("数据", "你的清洗算不算真正清洗？", "算可复现的数据治理：schema、语言、轮次、非空、去重、结构错误、风险评分、分层选择和人工审计。但不等于逐条事实核验；报告明确限制了这个结论。", "诚实边界比夸大“清洗1万条”更专业。"),
    ("数据", "100条审核够吗？", "用于工程门禁和发现明显系统性问题够用，但不能证明1万条全部正确。99/100意味着抽样观察到1% major，样本量小且单人标注；论文级别需要更大样本、双人标注和置信区间。", "可以给出局限，不要声称99%总体准确率。"),
    ("数据", "为什么clean/raw不使用相同prompt？", "当前上游没有同题高低质量答案对，所以采用任务桶和长度桶计数匹配，并保证问题集合不重叠。它减少主要结构混淆，但不能消除隐藏难度，是实验局限。", "更强设计是同prompt多答案或倾向评分匹配。"),
    ("数据", "NFKC去重做了什么？", "把兼容字符规范化，例如全角和半角形式，再转小写并去非字母数字字符，抓住表面格式不同的精确问题。它不能识别同义改写。", "语义去重可用embedding+ANN，但阈值和成本需验证。"),
    ("数据", "蓄水池抽样为什么均匀？", "处理第n个合格样本时，以k/n概率保留进容量k的池，因此归纳可证明任一已见样本最终留在池中的概率都是k/n。", "空间O(k)，适合流式7M。"),
    ("数据", "质量分公式是不是拍脑袋？", "权重属于透明启发式，不是理论最优。我把它定位为实验构造规则，并通过分布、审核和外部模型评测验证结果。要发论文需做权重敏感性分析或学习排序器。", "主动承认比伪装成科学真值更好。"),
    ("预处理", "为什么要messages/role/content？", "这是chat template消费的结构化对话接口。原始字段叫什么不重要，必须映射成目标tokenizer模板能够读取的role/content语义。不同框架也可能接受prompt/response，但最终仍要适配模型模板。", "ModelScope与Hugging Face是来源/接口差异，不决定对话schema。"),
    ("预处理", "apply_chat_template会自动生成回答吗？", "不会。它只是把结构化messages拼成模型需要的控制token序列。add_generation_prompt只补assistant开始标记，真正生成发生在model.generate。", "训练full不补生成提示，prompt分支补它用于定位答案起点。"),
    ("预处理", "为什么两次apply_chat_template？", "full得到问题+标准答案的完整token；prompt得到回答开始前的token前缀。两者相减才能精确构造仅监督assistant的labels。", "先检查prompt确实是full前缀，防止模板差异造成错位。"),
    ("预处理", "-100是什么意思？", "Transformers因果LM内部交叉熵使用ignore_index=-100；这些位置不计loss。模型仍读取prompt的input_ids作为条件，只是不被训练去复述问题。", "padding位置的labels同样补-100。"),
    ("预处理", "为什么不把所有token都算loss？", "那会让模型学习预测用户输入和模板，而目标是条件于用户输入生成助手答案；还会让长prompt主导loss。", "某些继续预训练场景会对全部token监督，但不是当前SFT目标。"),
    ("预处理", "为什么过滤而不是截断？", "右截断很可能切断assistant答案和结束token，留下残缺监督；左截断会丢问题条件。过滤保持样本语义完整，代价是损失部分长样本。", "若长上下文是目标，应提高max length或设计按消息边界的截断策略。"),
    ("预处理", "为什么长度1024？", "它是覆盖率、显存和速度的折中。attention成本近似O(L²)，2048相较1024注意力计算约4倍；当前项目不是长上下文训练。最终报告实际过滤比例而不是声称无损。", "1024不是模型最大能力，只是本实验预算。"),
    ("训练", "QLoRA量化的是哪些东西？", "冻结Base权重以4-bit存储；矩阵计算反量化到bf16/fp16。LoRA A/B保持可训练高精度，梯度只针对LoRA；优化器状态再用8-bit降低显存。", "4-bit Base仍参与前向和梯度对输入/LoRA的传播，但其参数不更新。"),
    ("训练", "NF4有什么特殊？", "NF4是为近似正态分布权重设计的非均匀4-bit码本，在有限16个取值下比均匀INT4更适配神经网络权重分布。", "double quant进一步量化每组量化尺度。"),
    ("训练", "LoRA为什么有效？", "假设任务适配所需的权重更新位于低维子空间，用BA近似完整ΔW。参数从d×d降到r(d_in+d_out)，显著减少可训练参数和优化器状态。", "这是假设，不保证所有任务同样低秩。"),
    ("训练", "r和alpha分别控制什么？", "r控制低秩容量和参数量；alpha通过alpha/r缩放更新强度。当前r=16、alpha=32，缩放为2。增大r不一定更好，还会提高显存和过拟合风险。", "不同target公平比较时最好固定总可训练参数预算。"),
    ("训练", "q/k/v/o怎么解释？", "Q表示当前token提出的查询，K表示候选token的索引特征，QK相似度决定注意力权重，V提供被聚合内容，O把多头结果映射回隐藏空间。", "LoRA挂在这些投影上能直接调整信息路由。"),
    ("训练", "gate/up/down是什么？", "up把隐藏维度扩展到中间维度，gate控制激活通道，down投影回隐藏维度。它们承载大量非线性特征变换，因此all-linear容量更大。", "更大覆盖也意味着更多LoRA参数和训练成本。"),
    ("训练", "覆盖82%为什么不是训练82%？", "82%描述目标线性层覆盖的原始参数体量，但原始W全部冻结；每层只新增低秩A/B。S2新增约3303万参数，占原模型约0.74%。", "必须区分覆盖模块、可训练参数和计算量。"),
    ("训练", "gradient accumulation等于大batch吗？", "累积8个micro-batch后再optimizer.step，使平均梯度近似batch16。区别包括数值舍入、dropout随机性、梯度裁剪时机以及每步调度定义。", "日志中的step通常指optimizer step。"),
    ("训练", "gradient checkpointing做了什么？", "不保存部分中间激活，反向时重新执行相关前向计算，因此用计算时间换激活显存。它不等于模型checkpoint文件。", "use_cache需关闭以避免与训练重算冲突。"),
    ("训练", "为什么是2e-4？", "LoRA常用学习率高于全参微调，因为只更新随机初始化的小矩阵；2e-4是经验起点。真正严谨需要小规模lr sweep，但消融中必须固定。", "不要声称它是理论最优。"),
    ("训练", "warmup_ratio=0.05有什么作用？", "训练初期LoRA和优化器状态尚未稳定，线性升高学习率能避免大步更新破坏表示；之后进入线性衰减。", "200步pilot约前10步warmup。"),
    ("训练", "为什么S1/S2显示1.76 epoch？", "配置同时有epoch和max_steps，但Transformers中正max_steps优先。200个optimizer step对应约1.76次遍历当前1815条、有效batch16的数据。", "这是预期行为，不是训练提前停止。"),
    ("训练", "eval loss为什么不能选胜者？", "它是teacher forcing下标准答案token的平均负对数似然，不能直接测自由生成时是否满足多条约束。S1/S2 eval loss只差约0.0012，需要外部Multi-IF。", "loss仍可用于发现发散、过拟合和选择同实验checkpoint。"),
    ("评测", "strict prompt和instruction有什么区别？", "prompt级要求一道题的全部约束都通过；instruction级把每条约束独立计分。一个三约束问题通过两条时，prompt=0、instruction=2/3。", "因此prompt更接近完整任务成功率。"),
    ("评测", "strict和loose有什么区别？", "strict直接检查原回答；loose允许移除首尾行、星号等有限格式包裹后重判。它用来区分核心内容失败和轻微包装偏差。", "loose仍是规则评测，不是LLM语义裁判。"),
    ("评测", "为什么turn2/3快速下降？", "后续轮次通常累积更多约束，而且评测把前轮模型回答放回上下文，早期错误会传播；完整prompt全约束通过的概率也随约束数增加而下降。", "需要结合instruction级指标区分局部能力与完整成功。"),
    ("评测", "为什么max_new_tokens=512？", "防止无限生成并固定成本，同时覆盖大多数约束回答。若触顶截断，会系统性压低长度/结尾相关指标，所以错误分析必须统计generated_tokens==512。", "所有实验必须相同，不能给某模型更长预算。"),
    ("评测", "为什么do_sample=false？", "贪心/确定性生成降低随机方差，让同一输入、模型和环境更容易复现。代价是无法评价多样性，也可能不是模型的最优采样表现。", "对控制实验优先复现性。"),
    ("评测", "Multi-IF能证明通用能力提升吗？", "不能。它主要测多约束和多轮指令遵循。还应加入知识、数学、代码或至少回归集，检查SFT是否造成灾难性遗忘。", "结论必须限定到评测覆盖范围。"),
    ("实验", "S1和S2是否严格公平？", "数据、step、batch、lr等一致，但可训练参数量和每步计算不同，所以它回答的是实际策略在固定step预算下的效果，不是固定参数预算下的模块纯贡献。", "论文级改进应调rank使总可训练参数相等。"),
    ("实验", "S3与clean胜者如何归因？", "固定target、初始化Base、训练步数和全部超参数，只改变质量分层的数据；clean/raw任务和长度桶计数一致。差异支持数据质量作用，但仍受隐藏难度影响。", "不能把相关结果夸成完全因果证明。"),
    ("实验", "S4为什么要重新从Base训练？", "S4研究clean数据规模从2k到10k；若接S1继续训，同时改变初始化和累计步数，无法判断提升来自数据规模还是继续优化。", "S4与pilot独立输出目录。"),
    ("工程", "final_adapter里为什么没有完整4B模型？", "LoRA只保存低秩增量、配置和tokenizer信息。推理时先加载同revision的Base，再通过PeftModel挂adapter。", "这也是adapter体积远小于基座的原因。"),
    ("工程", "--resume如何保证不断点重复？", "评测JSONL每完成一个row立即flush；重启时解析已有完整JSON行，按稳定id放入completed集合并跳过。末尾损坏行会忽略。", "它恢复到题级，不恢复一题内部的turn。"),
    ("工程", "你遇到过什么真实故障？", "AutoDL曾缺punkt_tab和pycrfsuite导致长评测中断，因此补齐NLTK资源和python-crfsuite，并实现--resume。还处理了HF网络不可达和本地离线缓存。", "回答要强调定位方法和防复发，而不是只说pip install。"),
    ("结果", "如果S2只比S1高0.5%，怎么选？", "同时考虑各turn、strict prompt、运行成本和误差稳定性。没有重复seed时，0.5%可能只是样本波动；若收益小而S2参数/时间明显更高，会倾向attention并把不确定性写清楚。", "不能机械选小数点更高者。"),
    ("结果", "项目最严重的限制是什么？", "数据质量分是启发式且只做精确污染检查；实验只有单seed；S1/S2参数预算不等；Multi-IF覆盖面窄。当前结论应定位为工程证据而非算法普遍规律。", "主动给出下一步：语义去重、多seed、等参数LoRA、通用能力回归。"),
    ("DPO", "为什么SFT后还要DPO？", "SFT学习参考答案分布；DPO利用chosen/rejected直接优化偏好差异，可针对格式、帮助性或约束遵循做进一步对齐。它不能弥补错误偏好数据。", "项目计划中DPO是SFT闭环后的独立阶段，不混入当前消融。"),
    ("DPO", "DPO和GRPO有什么区别？", "DPO是离线偏好优化，不需要在线rollout或显式reward；GRPO对同一prompt采样多条输出，按组内相对奖励进行策略更新，适合可自动验证任务，但计算和奖励设计更复杂。", "调用现成Trainer不等于算法创新。"),
]


def add_question_bank(doc):
    doc.add_page_break()
    doc.add_heading("第4章：30分钟以上项目深挖题库", level=1)
    add_note(
        doc,
        "使用方式",
        "先遮住答案口述60-90秒，再核对标准答案。面试官通常从一个薄弱回答连续追问，而不是按顺序问完。",
        LIGHT_BLUE,
    )
    category = None
    number = 0
    for item_category, question, answer, followup in QUESTIONS:
        if item_category != category:
            category = item_category
            doc.add_heading(f"{category}深挖", level=2)
        number += 1
        p = doc.add_paragraph()
        p.paragraph_format.keep_with_next = True
        set_run_font(p.add_run(f"Q{number}. {question}"), size=10.5, bold=True, color=DARK_BLUE)
        p = doc.add_paragraph(style="QAAnswer")
        set_run_font(p.add_run("标准回答："), size=9.8, bold=True, color=BLUE)
        set_run_font(p.add_run(answer), size=9.8)
        p = doc.add_paragraph(style="QAAnswer")
        set_run_font(p.add_run("继续追问/边界："), size=9.5, bold=True, color=RGBColor(122, 90, 0))
        set_run_font(p.add_run(followup), size=9.5)


def add_mock_interview(doc):
    doc.add_page_break()
    doc.add_heading("第5章：真实30分钟面试节奏", level=1)
    sections = [
        ("0-3分钟", "项目总述", "用一句定位、三组研究问题、当前结果边界说清项目。不要从安装环境讲起。"),
        ("3-9分钟", "数据治理", "重点解释硬删除/软风险、质量分、蓄水池抽样、分层匹配和100条审核边界。"),
        ("9-14分钟", "模板与labels", "现场画出prompt/full序列，解释-100、im_end、动态padding和1024过滤。"),
        ("14-21分钟", "QLoRA训练", "从W'=W+(alpha/r)BA讲到NF4、double quant、梯度累积、checkpointing和target模块。"),
        ("21-27分钟", "消融与评测", "证明每组只改变什么；解释strict/loose、prompt/instruction和多轮退化。"),
        ("27-30分钟", "结果与不足", "给出现有数字，不夸大；主动承认单seed、评分启发式、参数预算不等和评测覆盖窄。"),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    set_table_widths(table, [1.0, 1.25, 4.25])
    for i, text in enumerate(("时间", "主题", "你必须讲清楚")):
        shade_cell(table.rows[0].cells[i], LIGHT_BLUE)
        p = table.rows[0].cells[i].paragraphs[0]
        set_run_font(p.add_run(text), size=9.2, bold=True, color=DARK_BLUE)
    for row in sections:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_run_font(cells[i].paragraphs[0].add_run(value), size=9)

    doc.add_heading("三分钟开场稿", level=2)
    opening = (
        "这个项目不是简单跑LoRA，而是研究中文后训练中三个容易混在一起的因素：数据质量、数据规模和LoRA挂载范围。"
        "我以Qwen3-4B-Base为统一起点，从Infinity-Instruct 7M构建一问一答中文候选，通过技术过滤、NFKC精确去重、"
        "硬错误删除、软风险评分和任务/长度分层形成full clean 10k、clean 2k和raw 2k。clean/raw的问题不重叠，但分层计数一致。"
        "预处理使用Qwen chat template，prompt token的labels设为-100，只监督assistant答案并保留im_end；超长样本成对过滤。"
        "训练采用4-bit NF4 QLoRA，S1只挂q/k/v/o，S2挂全部Transformer线性层，固定200个optimizer steps和其他参数。"
        "随后用454条Multi-IF中文多轮任务，在固定greedy、512 token、thinking关闭的条件下比较外部指令遵循，而不是凭eval loss选模型。"
        "B0三轮official average为40.45%、20.46%、16.07%；S1/S2训练已完成，正式评测完成后选择target，再训练raw 2k的S3和full 10k的S4。"
        "我对结论的限制是：质量分是启发式、人工审核只有100条、消融目前是单seed，而且Multi-IF不代表全部通用能力。"
    )
    doc.add_paragraph(opening)

    doc.add_heading("面试时不要说的表述", level=2)
    add_bullets(
        doc,
        [
            "错误：我人工清洗了1万条。正确：我自动治理并冻结1万条，对其中100条做分层人工审计。",
            "错误：all-linear训练了原模型82%的参数。正确：覆盖这些线性层，但只训练新增低秩矩阵，原权重冻结。",
            "错误：eval loss更低所以模型更好。正确：loss只做训练诊断，最终看冻结外部生成评测。",
            "错误：Multi-IF证明模型全能力提升。正确：它支持中文多轮约束遵循范围内的结论。",
            "错误：负reward都是垃圾。正确：reward是上游弱信号，结合可观察风险用于排序。",
        ],
    )

    doc.add_heading("公开GitHub前检查清单", level=2)
    add_bullets(
        doc,
        [
            "补requirements或environment文件，锁定torch/transformers/peft/bitsandbytes等关键版本。",
            "记录third_party/Multi-IF的git commit；当前summary中的official_rule_revision为null。",
            "加入S1-S4最终指标表、可训练参数、训练时间、峰值显存和推理速度。",
            "保留answers JSONL用于错误分析，但不要提交模型缓存、密钥和绝对个人路径。",
            "README给出从数据门禁到评测的最短复现路径，并明确局限。",
        ],
    )


def add_appendix(doc):
    doc.add_page_break()
    doc.add_heading("附录：文件职责速查", level=1)
    items = [
        ("configs/project.yaml", "冻结所有实验参数"),
        ("src/build_sft_datasets.py", "构建full/clean/raw数据"),
        ("src/data_schema.py", "SFT/DPO schema验证"),
        ("scripts/validate_dataset.py", "原始和tokenized数据门禁"),
        ("scripts/audit_sft.py", "100条人工分层审核"),
        ("src/preprocess_sft.py", "切分、模板、tokenize、labels"),
        ("src/train_sft.py", "S1-S4 QLoRA训练"),
        ("scripts/prepare_multi_if.py", "冻结中文正式评测集"),
        ("src/evaluate_instruction.py", "共享加载/生成和冒烟规则"),
        ("src/evaluate_multi_if.py", "正式多轮规则评测"),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    set_table_widths(table, [2.5, 4.0])
    for i, text in enumerate(("文件", "一句话职责")):
        shade_cell(table.rows[0].cells[i], LIGHT_BLUE)
        set_run_font(table.rows[0].cells[i].paragraphs[0].add_run(text), size=9.3, bold=True, color=DARK_BLUE)
    for path, purpose in items:
        cells = table.add_row().cells
        set_run_font(cells[0].paragraphs[0].add_run(path), name="Consolas", size=8.7)
        set_run_font(cells[1].paragraphs[0].add_run(purpose), size=9)

    doc.add_heading("最终复习标准", level=2)
    add_bullets(
        doc,
        [
            "不看代码，能画出一条样本从原始row到loss的完整数据流。",
            "能推导有效batch、解释200 steps与1.76 epoch的关系。",
            "能写出LoRA公式并区分4-bit基座、LoRA参数、梯度和优化器状态。",
            "能用一个三约束例子解释prompt/instruction两类准确率。",
            "能说出每组消融只改变什么，以及仍未控制什么。",
            "能引用实际数字，同时准确限定结论边界。",
        ],
    )


def main():
    doc = Document()
    configure_document(doc)
    add_title_page(doc)
    add_contents(doc)
    add_overview(doc)
    for chapter in file_guides():
        add_file_chapter(doc, *chapter)
    add_question_bank(doc)
    add_mock_interview(doc)
    add_appendix(doc)

    props = doc.core_properties
    props.title = "Qwen3中文后训练项目：面试深挖与代码阅读手册"
    props.subject = "数据治理、QLoRA SFT、消融实验与Multi-IF评测"
    props.author = ""
    props.keywords = "Qwen3, QLoRA, SFT, Multi-IF, 面试"
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    main()
